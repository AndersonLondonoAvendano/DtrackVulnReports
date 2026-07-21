"""T-052: Generador de reportes Word (.docx)."""
from __future__ import annotations

import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Inches, Pt, RGBColor

from vulntrack.domain.value_objects.severity import Severity
from vulntrack.infrastructure.reports.chart_builder import (
    ChartBuilder,
    EvolutionRow,
    ProjectRow,
    ReportData,
)

# ─── Constantes de estilo ────────────────────────────────────────────────────

_NAVY = RGBColor(0x1F, 0x38, 0x64)  # #1F3864
_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
_LIGHT_BLUE = RGBColor(0xDE, 0xEB, 0xF7)  # fondo de fila par

_SEV_RGB: dict[Severity, RGBColor] = {
    Severity.CRITICAL: RGBColor(0xC0, 0x00, 0x00),
    Severity.HIGH: RGBColor(0xFF, 0x00, 0x00),
    Severity.MEDIUM: RGBColor(0xFF, 0x99, 0x00),
    Severity.LOW: RGBColor(0xFF, 0xFF, 0x00),
    Severity.UNASSIGNED: RGBColor(0xD9, 0xD9, 0xD9),
}

_SEV_LABELS: dict[Severity, str] = {
    Severity.CRITICAL: "Crítica",
    Severity.HIGH: "Alta",
    Severity.MEDIUM: "Media",
    Severity.LOW: "Baja",
    Severity.UNASSIGNED: "Sin asignar",
}


# ─── Helpers ─────────────────────────────────────────────────────────────────


def _set_cell_bg(cell: object, rgb: RGBColor) -> None:  # type: ignore[type-arg]
    """Establece color de fondo de una celda de tabla."""
    from lxml import etree  # noqa: PLC0415

    tc = cell._tc  # type: ignore[attr-defined]
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _header_row(table: object, headers: list[str]) -> None:  # type: ignore[type-arg]
    """Pinta la primera fila de una tabla con fondo azul marino y texto blanco."""
    row = table.rows[0]  # type: ignore[attr-defined]
    for cell, header in zip(row.cells, headers):
        cell.text = header
        _set_cell_bg(cell, _NAVY)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0] if para.runs else para.add_run(header)
        run.text = header
        run.bold = True
        run.font.color.rgb = _WHITE
        run.font.size = Pt(9)


def _add_chart_image(doc: object, buf: io.BytesIO, width_inches: float = 5.5) -> None:  # type: ignore[type-arg]
    """Inserta imagen PNG en el documento."""
    doc.add_picture(buf, width=Inches(width_inches))  # type: ignore[attr-defined]
    last_para = doc.paragraphs[-1]  # type: ignore[attr-defined]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _section_heading(doc: object, text: str, level: int = 1) -> None:  # type: ignore[type-arg]
    """Agrega un título de sección con color azul marino."""
    p = doc.add_heading(text, level=level)  # type: ignore[attr-defined]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = _NAVY


# ─── Heatmap helpers ──────────────────────────────────────────────────────────


def _heatmap_row(
    table: object,  # type: ignore[type-arg]
    row_idx: int,
    values: list[str | int | float],
    severity_col_indices: list[int],
    row_data: list[int],
) -> None:
    """Rellena una fila de tabla con colores de heatmap en columnas de severidad."""
    row = table.rows[row_idx]  # type: ignore[attr-defined]
    for col_idx, (cell, value) in enumerate(zip(row.cells, values)):
        cell.text = str(value)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if para.runs:
            para.runs[0].font.size = Pt(9)
        if col_idx in severity_col_indices:
            sev_order = [
                Severity.CRITICAL, Severity.HIGH, Severity.MEDIUM,
                Severity.LOW, Severity.UNASSIGNED,
            ]
            local_idx = severity_col_indices.index(col_idx)
            if local_idx < len(sev_order) and local_idx < len(row_data):
                count = row_data[local_idx]
                if count > 0:
                    _set_cell_bg(cell, _SEV_RGB[sev_order[local_idx]])
                    if para.runs:
                        para.runs[0].font.bold = True
                        if sev_order[local_idx] in (Severity.CRITICAL, Severity.HIGH):
                            para.runs[0].font.color.rgb = _WHITE


# ─── DocxGenerator ───────────────────────────────────────────────────────────


class DocxGenerator:
    """Implementa ReportGenerator para generar reportes Word (.docx)."""

    def generate(self, data: ReportData) -> bytes:
        doc = Document()

        # Márgenes
        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        builder = ChartBuilder()

        self._add_header(doc, data)
        self._add_kpi_banner(doc, data)
        self._add_executive_paragraph(doc, data)
        self._add_estado_actual(doc, data, builder)
        self._add_nuevas_periodo(doc, data, builder)
        self._add_evolucion(doc, data, builder)
        self._add_quarterly_progress(doc, data, builder)
        self._add_prioritized_findings(doc, data)
        self._add_conclusions(doc, data)
        self._add_signature_footer(doc)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ── Sección 1: Encabezado ────────────────────────────────────────────────

    def _add_header(self, doc: Document, data: ReportData) -> None:
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("REPORTE DE VULNERABILIDADES")
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = _NAVY

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = subtitle.add_run(f"Período: {data.period_label}")
        run2.font.size = Pt(13)
        run2.font.color.rgb = _NAVY

        meta_lines = [
            f"Fecha de generación: {data.generated_at.strftime('%Y-%m-%d %H:%M')}",
            f"Desde: {data.date_from}  |  Hasta: {data.date_to}",
            f"Elaborado por: {data.author}",
            "Fuente: Dependency-Track",
        ]
        for line in meta_lines:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)

        doc.add_paragraph()

    # ── Sección 2: Banner KPIs ───────────────────────────────────────────────

    def _add_kpi_banner(self, doc: Document, data: ReportData) -> None:
        _section_heading(doc, "1. Resumen ejecutivo — KPIs del período")

        table = doc.add_table(rows=2, cols=4)
        table.style = "Table Grid"

        kpis = [
            ("Vulnerabilidades\nvigentes", str(data.total_vigentes), _NAVY),
            ("Nuevas en\nel período", str(data.total_nuevas), RGBColor(0xC0, 0x00, 0x00)),
            ("Tratadas en\nel período", str(data.total_tratadas), RGBColor(0x70, 0xAD, 0x47)),
            ("Risk Score\nportafolio", f"{data.risk_score_portfolio:.1f}", RGBColor(0xFF, 0x99, 0x00)),
        ]

        for col_idx, (label, value, color) in enumerate(kpis):
            # Fila 0: etiqueta
            label_cell = table.rows[0].cells[col_idx]
            label_cell.text = label
            label_para = label_cell.paragraphs[0]
            label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if label_para.runs:
                label_para.runs[0].font.size = Pt(9)
                label_para.runs[0].bold = True
            _set_cell_bg(label_cell, _NAVY)
            if label_para.runs:
                label_para.runs[0].font.color.rgb = _WHITE

            # Fila 1: valor
            value_cell = table.rows[1].cells[col_idx]
            value_cell.text = value
            value_para = value_cell.paragraphs[0]
            value_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if value_para.runs:
                value_para.runs[0].font.size = Pt(22)
                value_para.runs[0].bold = True
                value_para.runs[0].font.color.rgb = color

        doc.add_paragraph()

    # ── Sección 3: Párrafo ejecutivo ─────────────────────────────────────────

    def _add_executive_paragraph(self, doc: Document, data: ReportData) -> None:
        _section_heading(doc, "2. Análisis ejecutivo")

        criticas = data.portfolio_metrics.get(Severity.CRITICAL, 0)
        altas = data.portfolio_metrics.get(Severity.HIGH, 0)
        kev_count = len(data.kev_hits)

        lines = [
            f"Durante el período {data.period_label} ({data.date_from} — {data.date_to}) "
            f"el portafolio registró {data.total_vigentes} vulnerabilidades vigentes "
            f"({criticas} críticas, {altas} altas), con {data.total_nuevas} nuevas "
            f"identificaciones y {data.total_tratadas} vulnerabilidades tratadas.",
        ]
        if kev_count > 0:
            lines.append(
                f"Se detectaron {kev_count} vulnerabilidades en el catálogo CISA KEV, "
                "que requieren atención inmediata por existir exploits activos confirmados."
            )

        for line in lines:
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(4)

        doc.add_paragraph()

    # ── Sección 4: Estado actual ──────────────────────────────────────────────

    def _add_estado_actual(
        self, doc: Document, data: ReportData, builder: ChartBuilder
    ) -> None:
        _section_heading(doc, "3. Estado actual del portafolio")

        # Donut + barra
        donut_buf = builder.donut_by_severity(
            data.portfolio_metrics, "Distribución de vulnerabilidades vigentes"
        )
        _add_chart_image(doc, donut_buf, width_inches=4.5)

        barra_buf = builder.horizontal_bars_by_project(
            data.project_rows, "Vulnerabilidades vigentes por proyecto"
        )
        _add_chart_image(doc, barra_buf, width_inches=6.0)

        # Tabla heatmap
        p = doc.add_paragraph("Tabla de vulnerabilidades vigentes por proyecto:")
        p.paragraph_format.space_after = Pt(2)

        headers = ["Proyecto", "Crítica", "Alta", "Media", "Baja", "Sin asignar", "Total", "Risk Score"]
        table = doc.add_table(rows=len(data.project_rows) + 1, cols=len(headers))
        table.style = "Table Grid"
        _header_row(table, headers)

        for row_idx, pr in enumerate(data.project_rows, start=1):
            row_counts = [pr.critical, pr.high, pr.medium, pr.low, pr.unassigned]
            values: list[str | int | float] = [
                pr.name, pr.critical, pr.high, pr.medium, pr.low,
                pr.unassigned, pr.total, f"{pr.risk_score:.1f}",
            ]
            _heatmap_row(table, row_idx, values, [1, 2, 3, 4, 5], row_counts)

        doc.add_paragraph()

    # ── Sección 5: Nuevas en el período ──────────────────────────────────────

    def _add_nuevas_periodo(
        self, doc: Document, data: ReportData, builder: ChartBuilder
    ) -> None:
        _section_heading(doc, "4. Nuevas vulnerabilidades en el período")

        donut_buf = builder.donut_by_severity(
            data.new_portfolio_metrics, "Nuevas vulnerabilidades por severidad"
        )
        _add_chart_image(doc, donut_buf, width_inches=4.5)

        total_nuevas_p = doc.add_paragraph(
            f"Total de nuevas vulnerabilidades identificadas: {data.total_nuevas}"
        )
        total_nuevas_p.paragraph_format.space_after = Pt(4)

        # Tabla resumen nuevas por severidad
        headers = ["Severidad", "Cantidad"]
        table = doc.add_table(rows=len(Severity) + 1, cols=2)
        table.style = "Table Grid"
        _header_row(table, headers)

        sev_order = [
            Severity.CRITICAL, Severity.HIGH, Severity.MEDIUM,
            Severity.LOW, Severity.UNASSIGNED,
        ]
        for row_idx, sev in enumerate(sev_order, start=1):
            count = data.new_portfolio_metrics.get(sev, 0)
            row = table.rows[row_idx]
            row.cells[0].text = _SEV_LABELS[sev]
            row.cells[1].text = str(count)
            row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if count > 0:
                _set_cell_bg(row.cells[1], _SEV_RGB[sev])
                if row.cells[1].paragraphs[0].runs:
                    run = row.cells[1].paragraphs[0].runs[0]
                    run.bold = True
                    if sev in (Severity.CRITICAL, Severity.HIGH):
                        run.font.color.rgb = _WHITE

        doc.add_paragraph()

    # ── Sección 6: Evolución ──────────────────────────────────────────────────

    def _add_evolucion(
        self, doc: Document, data: ReportData, builder: ChartBuilder
    ) -> None:
        _section_heading(doc, "5. Evolución del período")

        divergent_buf = builder.divergent_bars_evolution(
            data.evolution_rows, "Variación de vulnerabilidades por proyecto"
        )
        _add_chart_image(doc, divergent_buf, width_inches=6.0)

        grouped_buf = builder.grouped_bars_inicio_vs_actual(
            data.evolution_rows, "Inicio vs Actual por proyecto"
        )
        _add_chart_image(doc, grouped_buf, width_inches=6.0)

        # Tabla evolución
        headers = ["Proyecto", "Inicio", "Actual", "Variación", "Tratadas"]
        table = doc.add_table(rows=len(data.evolution_rows) + 1, cols=len(headers))
        table.style = "Table Grid"
        _header_row(table, headers)

        for row_idx, er in enumerate(data.evolution_rows, start=1):
            row = table.rows[row_idx]
            values_ev = [er.name, er.inicio, er.actual, er.variacion, er.tratadas]
            for col_idx, val in enumerate(values_ev):
                cell = row.cells[col_idx]
                cell.text = (
                    f"{val:+d}" if col_idx == 3 and isinstance(val, int) else str(val)
                )
                para = cell.paragraphs[0]
                para.alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT if col_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
                )
                if para.runs:
                    para.runs[0].font.size = Pt(9)
                # Color variación: rojo si sube, verde si baja
                if col_idx == 3 and isinstance(val, int):
                    if val > 0:
                        _set_cell_bg(cell, RGBColor(0xFF, 0xCC, 0xCC))
                    elif val < 0:
                        _set_cell_bg(cell, RGBColor(0xCC, 0xFF, 0xCC))
                # Color tratadas: verde si > 0
                if col_idx == 4 and isinstance(val, int) and val > 0:
                    _set_cell_bg(cell, RGBColor(0x70, 0xAD, 0x47))
                    if para.runs:
                        para.runs[0].font.color.rgb = _WHITE

        doc.add_paragraph()

    # ── Sección 6: Avance de remediación por Q ────────────────────────────────

    def _add_quarterly_progress(
        self, doc: Document, data: ReportData, builder: ChartBuilder
    ) -> None:
        _section_heading(doc, "6. Avance de remediación por Q")

        qp = data.quarterly_progress
        if qp is None:
            doc.add_paragraph(
                "No hay sprints ni tratamientos registrados para este período."
            )
            doc.add_paragraph()
            return

        p = doc.add_paragraph(f"Período: {qp.anio}-Q{qp.trimestre}")
        p.paragraph_format.space_after = Pt(4)

        chart_buf = builder.quarterly_progress_bars(
            qp, f"Avance de remediación — {qp.anio}-Q{qp.trimestre}"
        )
        _add_chart_image(doc, chart_buf, width_inches=6.0)

        headers = [
            "Entraron", "Resueltas", "Pospuestas", "No cumplidas", "En curso",
            "Descartadas", "% cumplimiento",
        ]
        table = doc.add_table(rows=2, cols=len(headers))
        table.style = "Table Grid"
        _header_row(table, headers)

        pct_str = f"{qp.pct_cumplimiento * 100:.1f}%" if qp.pct_cumplimiento is not None else "—"
        values = [
            qp.entraron, qp.resueltas, qp.pospuestas, qp.no_cumplidas,
            qp.en_curso, qp.descartadas, pct_str,
        ]
        row = table.rows[1]
        for col_idx, val in enumerate(values):
            cell = row.cells[col_idx]
            cell.text = str(val)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(10)
                cell.paragraphs[0].runs[0].bold = True

        doc.add_paragraph()

        if qp.tendencia_por_sprint:
            p2 = doc.add_paragraph("Tendencia por sprint:")
            p2.paragraph_format.space_after = Pt(2)

            trend_headers = [
                "Sprint", "Entraron", "Resueltas", "Pospuestas",
                "No cumplidas", "En curso", "Descartadas",
            ]
            trend_table = doc.add_table(
                rows=len(qp.tendencia_por_sprint) + 1, cols=len(trend_headers)
            )
            trend_table.style = "Table Grid"
            _header_row(trend_table, trend_headers)

            for row_idx, t in enumerate(qp.tendencia_por_sprint, start=1):
                trend_row = trend_table.rows[row_idx]
                trend_values = [
                    t.sprint_nombre, t.entraron, t.resueltas, t.pospuestas,
                    t.no_cumplidas, t.en_curso, t.descartadas,
                ]
                for col_idx, val in enumerate(trend_values):
                    cell = trend_row.cells[col_idx]
                    cell.text = str(val)
                    para = cell.paragraphs[0]
                    para.alignment = (
                        WD_ALIGN_PARAGRAPH.LEFT if col_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
                    )
                    if para.runs:
                        para.runs[0].font.size = Pt(9)

        doc.add_paragraph()

    # ── Sección 7: Hallazgos priorizados ──────────────────────────────────────

    def _add_prioritized_findings(self, doc: Document, data: ReportData) -> None:
        _section_heading(doc, "7. Hallazgos priorizados")

        top_findings = data.prioritized_findings[:20]
        if not top_findings:
            doc.add_paragraph("No hay hallazgos priorizados en este período.")
            doc.add_paragraph()
            return

        headers = ["CVE / ID", "Componente", "Proyecto", "Severidad", "CVSS", "EPSS", "KEV", "Score"]
        table = doc.add_table(rows=len(top_findings) + 1, cols=len(headers))
        table.style = "Table Grid"
        _header_row(table, headers)

        for row_idx, fr in enumerate(top_findings, start=1):
            row = table.rows[row_idx]
            values = [
                fr.vuln_id,
                f"{fr.component_name} {fr.component_version or ''}".strip(),
                fr.project_name,
                _SEV_LABELS.get(fr.severity, fr.severity),
                f"{fr.cvss_v3_base_score:.1f}" if fr.cvss_v3_base_score is not None else "—",
                f"{fr.epss_score:.3f}" if fr.epss_score is not None else "—",
                "Sí" if fr.is_kev else "No",
                f"{fr.priority_score:.1f}",
            ]
            for col_idx, val in enumerate(values):
                cell = row.cells[col_idx]
                cell.text = str(val)
                para = cell.paragraphs[0]
                para.alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT if col_idx in (0, 1, 2) else WD_ALIGN_PARAGRAPH.CENTER
                )
                if para.runs:
                    para.runs[0].font.size = Pt(8)
            # Color severidad
            sev_cell = row.cells[3]
            _set_cell_bg(sev_cell, _SEV_RGB[fr.severity])
            if sev_cell.paragraphs[0].runs:
                sev_cell.paragraphs[0].runs[0].font.bold = True
                if fr.severity in (Severity.CRITICAL, Severity.HIGH):
                    sev_cell.paragraphs[0].runs[0].font.color.rgb = _WHITE
            # KEV badge
            if fr.is_kev:
                kev_cell = row.cells[6]
                _set_cell_bg(kev_cell, RGBColor(0xC0, 0x00, 0x00))
                if kev_cell.paragraphs[0].runs:
                    kev_cell.paragraphs[0].runs[0].font.color.rgb = _WHITE
                    kev_cell.paragraphs[0].runs[0].font.bold = True

        doc.add_paragraph()

    # ── Sección 8: Conclusiones ───────────────────────────────────────────────

    def _add_conclusions(self, doc: Document, data: ReportData) -> None:
        _section_heading(doc, "8. Conclusiones y recomendaciones")

        observations = self._build_observations(data)
        for obs in observations:
            p = doc.add_paragraph(obs, style="List Bullet")
            p.paragraph_format.space_after = Pt(3)

        doc.add_paragraph()

    def _build_observations(self, data: ReportData) -> list[str]:
        obs: list[str] = []

        criticas = data.portfolio_metrics.get(Severity.CRITICAL, 0)
        if criticas > 0:
            obs.append(
                f"Existen {criticas} vulnerabilidades críticas vigentes que requieren atención inmediata."
            )

        kev_count = len(data.kev_hits)
        if kev_count > 0:
            obs.append(
                f"Se identificaron {kev_count} vulnerabilidades en el catálogo CISA KEV con "
                "exploits activos confirmados. Se recomienda su remediación en un plazo máximo de 7 días."
            )

        if data.total_tratadas > 0:
            pct = (data.total_tratadas / max(data.total_vigentes + data.total_tratadas, 1)) * 100
            obs.append(
                f"Durante el período se trató el {pct:.0f}% de las vulnerabilidades detectadas "
                f"({data.total_tratadas} de {data.total_vigentes + data.total_tratadas})."
            )

        if not obs:
            obs.append(
                "No se registraron vulnerabilidades críticas en el período. Continuar con el monitoreo periódico."
            )

        return obs

    # ── Sección 9: Pie de firma ───────────────────────────────────────────────

    def _add_signature_footer(self, doc: Document) -> None:
        _section_heading(doc, "9. Firmas")

        table = doc.add_table(rows=2, cols=3)
        table.style = "Table Grid"

        labels = ["Elaboró", "Revisó", "Aprobó"]
        for col_idx, label in enumerate(labels):
            cell = table.rows[0].cells[col_idx]
            cell.text = label
            _set_cell_bg(cell, _NAVY)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if para.runs:
                para.runs[0].bold = True
                para.runs[0].font.color.rgb = _WHITE

        for col_idx in range(3):
            cell = table.rows[1].cells[col_idx]
            cell.text = ""
            # Espacio para firma manuscrita
            for _ in range(3):
                cell.add_paragraph()
