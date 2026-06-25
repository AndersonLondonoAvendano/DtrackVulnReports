"""Tests para T-052: DocxGenerator."""
from __future__ import annotations

import io
from datetime import date, datetime

from docx import Document

from vulntrack.domain.value_objects.priority_score import PriorityBand
from vulntrack.domain.value_objects.severity import Severity
from vulntrack.infrastructure.reports.chart_builder import (
    EvolutionRow,
    PrioritizedFindingRow,
    ProjectRow,
    ReportData,
)
from vulntrack.infrastructure.reports.docx_generator import DocxGenerator


def _sample_data() -> ReportData:
    return ReportData(
        period_label="Q2 2026",
        date_from=date(2026, 4, 1),
        date_to=date(2026, 6, 30),
        generated_at=datetime(2026, 6, 25, 10, 0, 0),
        author="Anderson Avendaño",
        total_vigentes=223,
        total_nuevas=118,
        total_tratadas=97,
        risk_score_portfolio=7.8,
        portfolio_metrics={
            Severity.CRITICAL: 12,
            Severity.HIGH: 45,
            Severity.MEDIUM: 98,
            Severity.LOW: 58,
            Severity.UNASSIGNED: 10,
        },
        new_portfolio_metrics={
            Severity.CRITICAL: 5,
            Severity.HIGH: 30,
            Severity.MEDIUM: 60,
            Severity.LOW: 23,
            Severity.UNASSIGNED: 0,
        },
        project_rows=[
            ProjectRow("daviplata-webview-frontend", 5, 10, 20, 15, 2, 52, 8.5),
            ProjectRow("backend-core", 2, 3, 8, 5, 0, 18, 4.2),
            ProjectRow("infra-base", 0, 0, 0, 0, 0, 0, 0.0),
        ],
        evolution_rows=[
            EvolutionRow("daviplata-webview-frontend", 60, 52, -8, 8),
            EvolutionRow("backend-core", 14, 18, 4, 0),
        ],
        prioritized_findings=[
            PrioritizedFindingRow(
                vuln_id="CVE-2024-1234",
                component_name="log4j-core",
                component_version="2.14.0",
                project_name="daviplata-webview-frontend",
                severity=Severity.CRITICAL,
                cvss_v3_base_score=9.8,
                epss_score=0.85,
                is_kev=True,
                priority_score=92.5,
                priority_band=PriorityBand.IMMEDIATE,
            ),
            PrioritizedFindingRow(
                vuln_id="CVE-2023-9999",
                component_name="spring-core",
                component_version="5.3.0",
                project_name="backend-core",
                severity=Severity.HIGH,
                cvss_v3_base_score=7.5,
                epss_score=0.42,
                is_kev=False,
                priority_score=65.0,
                priority_band=PriorityBand.HIGH,
            ),
        ],
        kev_hits=[
            PrioritizedFindingRow(
                vuln_id="CVE-2024-1234",
                component_name="log4j-core",
                component_version="2.14.0",
                project_name="daviplata-webview-frontend",
                severity=Severity.CRITICAL,
                cvss_v3_base_score=9.8,
                epss_score=0.85,
                is_kev=True,
                priority_score=92.5,
                priority_band=PriorityBand.IMMEDIATE,
            )
        ],
    )


class TestDocxGenerator:
    def setup_method(self) -> None:
        self.generator = DocxGenerator()
        self.data = _sample_data()

    def test_generate_returns_bytes(self) -> None:
        result = self.generator.generate(self.data)
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_generate_is_valid_docx(self) -> None:
        result = self.generator.generate(self.data)
        doc = Document(io.BytesIO(result))
        assert doc is not None

    def test_docx_has_content(self) -> None:
        result = self.generator.generate(self.data)
        doc = Document(io.BytesIO(result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Q2 2026" in full_text
        assert "Anderson Avendaño" in full_text

    def test_docx_contains_kpis(self) -> None:
        result = self.generator.generate(self.data)
        doc = Document(io.BytesIO(result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "223" in full_text  # total_vigentes
        assert "118" in full_text  # total_nuevas
        assert "97" in full_text   # total_tratadas

    def test_docx_has_tables(self) -> None:
        result = self.generator.generate(self.data)
        doc = Document(io.BytesIO(result))
        assert len(doc.tables) >= 3

    def test_generate_with_empty_data(self) -> None:
        from dataclasses import replace
        empty_data = ReportData(
            period_label="Q1 2026",
            date_from=date(2026, 1, 1),
            date_to=date(2026, 3, 31),
            generated_at=datetime(2026, 3, 31, 0, 0, 0),
            author="Test",
            total_vigentes=0,
            total_nuevas=0,
            total_tratadas=0,
            risk_score_portfolio=0.0,
        )
        result = self.generator.generate(empty_data)
        assert isinstance(result, bytes)
        assert len(result) > 0
        doc = Document(io.BytesIO(result))
        assert doc is not None

    def test_generate_with_no_kev_hits(self) -> None:
        from dataclasses import replace
        data_no_kev = ReportData(
            period_label="Q2 2026",
            date_from=date(2026, 4, 1),
            date_to=date(2026, 6, 30),
            generated_at=datetime(2026, 6, 25, 10, 0, 0),
            author="Test",
            total_vigentes=10,
            total_nuevas=5,
            total_tratadas=3,
            risk_score_portfolio=3.5,
            portfolio_metrics={Severity.HIGH: 10},
            kev_hits=[],
        )
        result = self.generator.generate(data_no_kev)
        assert len(result) > 0
